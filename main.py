from flask import Flask, request, jsonify, send_from_directory, render_template
import urllib.parse
import requests
import os
import re
import sqlite3
import tempfile
from docx import Document
from datetime import datetime

app = Flask(__name__, template_folder='template')

# Gunakan direktori sementara untuk Vercel atau lingkungan lain yang read-only
TEMP_DIR = tempfile.gettempdir()

# Batas maksimum panjang prompt (dalam karakter)
MAX_PROMPT_LENGTH = 1500
DATABASE = os.path.join(TEMP_DIR, 'novels.db')  # Database disimpan di /tmp/

def init_db():
    conn = sqlite3.connect(DATABASE)
    c = conn.cursor()
    c.execute('''
        CREATE TABLE IF NOT EXISTS chapters (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            novel_title TEXT,
            chapter TEXT,
            chapter_title TEXT,
            chapter_order INTEGER,
            narrative_type TEXT,
            content TEXT,
            doc_filepath TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    conn.commit()
    conn.close()

# Pastikan database diinisialisasi sebelum request pertama
db_initialized = False
@app.before_request
def initialize_db_once():
    global db_initialized
    if not db_initialized:
        init_db()
        db_initialized = True

def delete_old_chapters():
    conn = sqlite3.connect(DATABASE)
    c = conn.cursor()
    c.execute("SELECT id, doc_filepath FROM chapters WHERE created_at < datetime('now', '-7 days')")
    rows = c.fetchall()
    for row in rows:
        _, doc_filepath = row
        # doc_filepath adalah relative path; buat full path dari TEMP_DIR
        full_path = os.path.join(TEMP_DIR, doc_filepath)
        if doc_filepath and os.path.exists(full_path):
            try:
                os.remove(full_path)
            except Exception as e:
                print("Gagal menghapus file:", full_path, e)
    c.execute("DELETE FROM chapters WHERE created_at < datetime('now', '-7 days')")
    conn.commit()
    conn.close()

def sanitize_title(title):
    return re.sub(r'\W+', '', title.replace(" ", "_"))

def get_chapter_order(chapter_input):
    chapter_lower = chapter_input.lower().strip()
    if chapter_lower == "prolog":
        return 0
    else:
        match = re.search(r'bab\s*(\d+)', chapter_lower)
        return int(match.group(1)) if match else None

def get_context_from_db(novel_title, new_order):
    conn = sqlite3.connect(DATABASE)
    c = conn.cursor()
    c.execute("""
        SELECT chapter, content FROM chapters 
        WHERE novel_title = ? AND chapter_order < ? 
        ORDER BY chapter_order ASC
    """, (novel_title, new_order))
    rows = c.fetchall()
    conn.close()
    context = ""
    for chapter, content in rows:
        context += f"{chapter} content: {content}\n"
    return context

def add_markdown_line_to_paragraph(paragraph, text):
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)

def add_markdown_to_doc(doc, markdown_text):
    lines = markdown_text.splitlines()
    for line in lines:
        line = line.strip()
        if not line:
            continue
        if line.startswith("#"):
            level = len(line) - len(line.lstrip("#"))
            text = line.lstrip("#").strip()
            level = level if level <= 4 else 4
            p = doc.add_heading("", level=level)
            add_markdown_line_to_paragraph(p, text)
        elif line.startswith("- "):
            p = doc.add_paragraph(style='List Bullet')
            add_markdown_line_to_paragraph(p, line[2:])
        else:
            p = doc.add_paragraph()
            add_markdown_line_to_paragraph(p, line)

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/generate_novel', methods=['POST'])
def generate_novel_endpoint():
    try:
        delete_old_chapters()
        data = request.json

        chapter_input   = data.get("chapter", "Prolog")
        character_name  = data.get("character_name", "Alex")
        genre           = data.get("genre", "Fantasy")
        world_setting   = data.get("world_setting", "Dunia paralel penuh keajaiban")
        conflict        = data.get("conflict", "Menyelamatkan dunia dari ancaman gelap")
        special_power   = data.get("special_power", "Mengendalikan elemen")
        plot_twist      = data.get("plot_twist", "Musuh utama ternyata saudara kembarnya")
        writing_style   = data.get("writing_style", "Misterius dan dramatis")
        narrative_type  = data.get("narrative_type", "linear")
        novel_title     = data.get("novel_title", "Novel Tanpa Judul")
        chapter_title   = data.get("chapter_title", "")
        chapter_instructions = data.get("chapter_instructions", 
            "Ceritakan bab ini dengan detail, sertakan dialog antar karakter dan narasi yang hidup.")

        new_order = get_chapter_order(chapter_input)
        if new_order is None:
            return jsonify({"status": "error", "message": "Format chapter tidak valid. Gunakan 'Prolog' atau 'Bab <nomor>'."}), 400

        # Buat folder relatif (misalnya: novel_KisahPohonMisterius) di TEMP_DIR
        relative_folder = f"novel_{sanitize_title(novel_title)}"
        folder_name = os.path.join(TEMP_DIR, relative_folder)
        if not os.path.exists(folder_name):
            os.makedirs(folder_name)

        context_prompt = get_context_from_db(novel_title, new_order) if narrative_type.lower() == "linear" else ""

        if chapter_input.lower() == "prolog":
            chapter_prompt = f"""
            === {chapter_input.upper()} ===
            Tuliskan prolog dengan pengenalan dunia dan karakter secara mendalam.
            Gunakan deskripsi, dialog, dan narasi untuk memperkenalkan latar cerita.
            Informasi tambahan:\n- Genre: {genre}\n- Dunia: {world_setting}\n- Tokoh Utama: {character_name} dengan kekuatan {special_power}\n- Konflik: {conflict}\n- Plot Twist: {plot_twist}\n- Gaya: {writing_style}\n\n{chapter_instructions} dengan prolog yang benar
            """
        else:
            chapter_prompt = f"""
            === {chapter_input.upper()} ===
            Tuliskan bab ini sebagai kelanjutan cerita yang naratif, dengan dialog antar karakter, deskripsi mendalam, dan alur cerita yang koheren.
            Jangan hanya menampilkan template, tetapi kembangkan cerita menjadi narasi yang hidup.
            Gunakan informasi berikut sebagai dasar:\nGenre: {genre}\nDunia: {world_setting}\nTokoh Utama: {character_name} dengan kekuatan {special_power}\nKonflik: {conflict}\nPlot Twist: {plot_twist}\nGaya: {writing_style}\n\n{chapter_instructions} dengan cerita yang benar
            """
        # Hapus bagian template_info dari prompt, gunakan hanya chapter_prompt dan context_prompt
        full_prompt = chapter_prompt + "\n" + context_prompt
        if len(full_prompt) > MAX_PROMPT_LENGTH:
            required_prompt = chapter_prompt + "\n"
            allowed_context_length = MAX_PROMPT_LENGTH - len(required_prompt)
            allowed_context_length = allowed_context_length if allowed_context_length > 0 else 0
            trimmed_context = context_prompt[-allowed_context_length:] if allowed_context_length > 0 else ""
            full_prompt = required_prompt + trimmed_context

        encoded_prompt = urllib.parse.quote(full_prompt)
        pollinations_url = f"https://text.pollinations.ai/openai/{encoded_prompt}"
        response = requests.get(pollinations_url)
        if response.status_code == 200:
            generated_story = response.text

            # Buat dokumen Word dan simpan ke folder yang sudah dibuat
            doc = Document()
            doc.add_heading(chapter_title if chapter_title else chapter_input, level=1)
            add_markdown_to_doc(doc, generated_story)
            doc_filename = f"prolog.doc - {sanitize_title(novel_title)}.doc" if chapter_input.lower() == "prolog" else f"bab{new_order} - {sanitize_title(novel_title)}.doc"
            doc_filepath = os.path.join(folder_name, doc_filename)
            doc.save(doc_filepath)
            
            # Buat relative path (tanpa TEMP_DIR) untuk disimpan ke DB dan digunakan di endpoint download
            relative_doc_path = os.path.join(relative_folder, doc_filename)
            
            conn = sqlite3.connect(DATABASE)
            c = conn.cursor()
            c.execute("""
                INSERT INTO chapters (novel_title, chapter, chapter_title, chapter_order, narrative_type, content, doc_filepath)
                VALUES (?, ?, ?, ?, ?, ?, ?)
            """, (novel_title, chapter_input, chapter_title, new_order, narrative_type, generated_story, relative_doc_path))
            conn.commit()
            conn.close()

            return jsonify({
                "status": "success",
                "novel_title": novel_title,
                "chapter": chapter_input,
                "order": new_order,
                "doc_filepath": relative_doc_path,
                "novel": generated_story
            })
        else:
            return jsonify({"status": "error", "message": "Gagal mengambil cerita dari AI"}), 500
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500

@app.route('/download/<path:filename>', methods=['GET'])
def download_file(filename):
    try:
        # filename diharapkan berupa relative path, misalnya: novel_KisahPohonMisterius/bab1.doc
        return send_from_directory(TEMP_DIR, filename, as_attachment=True)
    except Exception as e:
        return str(e), 404

if __name__ == '__main__':
    init_db()
    app.run(debug=True)
